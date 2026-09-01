"""Layer 6 — assembler, cache, and the FastAPI app.

All tests are offline: S3 mocked with moto, DB in-memory, no LibreOffice.

The assembler tests run against the PRISTINE template in the resume guide, not
against the operator's personalised copy — that copy is gitignored (it carries
real contact details and hyperlinks), so a test depending on it would silently
skip in CI and on any other machine. Structure is identical between the two;
only the text differs, and the assembler matches on structure alone.
"""

from __future__ import annotations

import json
import re
import shutil
import zipfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

REPO_ROOT = Path(__file__).resolve().parent.parent
#: The operator's personalised template (gitignored, may be absent).
TEMPLATE_PATH = REPO_ROOT / "resumes" / "templates" / "headless_v1.docx"
#: The pristine upstream template — present wherever the guide is checked out.
PRISTINE_PATH = Path(
    "/home/vishnu/projects/resume guide/Headless+Resume+Template.docx"
)


def _template() -> Path:
    for candidate in (TEMPLATE_PATH, PRISTINE_PATH):
        if candidate.exists():
            return candidate
    pytest.skip("no Headless template available")


# ---------------------------------------------------------------------------
# Assembler
# ---------------------------------------------------------------------------


@pytest.fixture
def minimal_profile(tmp_path) -> Path:
    data = {
        "work_experience": [{
            "id": "exp1",
            "role_blocks": [{
                "role": "backend",
                "entry_header": "Backend Engineer at TechCorp, Pune",
                "entry_dates": "January 2024 to current",
                "title_aliases": ["Backend Engineer"],
                "bullets": [
                    {"id": "b1", "text": "Kept the platform running for customers."},
                    {"id": "b2", "text": "Built APIs serving 10k requests."},
                ],
                "extra_bullets": [
                    {"id": "x1", "text": "Ran services in Docker to cut setup time."},
                ],
            }],
        }],
        "projects": [{
            "id": "proj1", "name": "Stock Prediction Engine",
            "link": "https://github.com/user/stock",
            "role_blocks": [{
                "role": "ml",
                "entry_header": "Stock Prediction Engine",
                "title_aliases": ["Machine Learning Engineer"],
                "bullets": [
                    {"id": "pb1", "text": "Predicted next-minute price direction."},
                    {"id": "pb2", "text": "Modelled with scikit-learn."},
                ],
            }],
        }],
    }
    path = tmp_path / "master_profile.json"
    path.write_text(json.dumps(data))
    return path


@pytest.fixture
def minimal_selection():
    from src.llm.schemas import SelectedEntryOut, StoredSelection

    return StoredSelection(
        job_id="job001",
        template_version="abc12345",
        built_at="2026-08-31T00:00:00Z",
        jd_keywords=["Python", "Docker"],
        keyword_coverage=0.6,
        lead_entry_coverage=0.5,
        entries=[
            SelectedEntryOut(
                kind="work", entry_id="exp1", block_id="exp1::backend",
                title_alias="Backend Engineer",
                header_left="Backend Engineer at TechCorp, Pune",
                header_right="January 2024 to current",
                bullet_ids=["b1", "b2", "x1"], covered=["Docker"], cap=6,
            ),
            SelectedEntryOut(
                kind="project", entry_id="proj1", block_id="proj1::ml",
                title_alias="Machine Learning Engineer",
                header_left="Stock Prediction Engine",
                header_right="https://github.com/user/stock",
                bullet_ids=["pb1", "pb2"], cap=5,
            ),
        ],
    )


def _assemble(profile, selection, tmp_path):
    from src.endpoint.assembler import assemble_docx

    out = tmp_path / "out.docx"
    assemble_docx(selection, profile, _template(), out)
    return Document(str(out))


def test_assemble_docx_renders_both_sections(minimal_profile, minimal_selection, tmp_path):
    doc = _assemble(minimal_profile, minimal_selection, tmp_path)
    text = [p.text for p in doc.paragraphs]

    # The template ships ONE section heading; the assembler mints the second.
    assert "Work History" in text
    assert "Projects" in text
    assert text.index("Work History") < text.index("Projects")

    assert "Backend Engineer at TechCorp, Pune\tJanuary 2024 to current" in text
    assert "Built APIs serving 10k requests." in text
    assert "Ran services in Docker to cut setup time." in text, (
        "an extra_bullet selected by the greedy must render like any other"
    )
    assert "Stock Prediction Engine\tCode \u2192" in text


def test_a_project_entry_shows_a_short_hyperlink_where_a_job_shows_dates(
    minimal_profile, minimal_selection, tmp_path
):
    """The full URL overflows the line — measured at 112 chars against an ~89-char
    budget at Arial 10.5 across 6.5in — so the slot carries a short link instead."""
    doc = _assemble(minimal_profile, minimal_selection, tmp_path)
    line = next(p for p in doc.paragraphs if p.text.startswith("Stock Prediction"))
    assert line.text.endswith("Code \u2192")
    assert len(line.text.replace("\t", "")) < 89

    rid = line._p.find(qn("w:hyperlink")).get(qn("r:id"))
    assert doc.part.rels[rid].target_ref == "https://github.com/user/stock"


def test_the_minted_hyperlink_has_no_dangling_relationship(
    minimal_profile, minimal_selection, tmp_path
):
    """A dangling r:id makes LibreOffice refuse the file, which fails the PDF
    render rather than merely losing a link. This is why the relationship is
    minted through python-docx's relate_to() instead of by patching the saved zip.
    """
    from src.endpoint.assembler import assemble_docx

    out = tmp_path / "out.docx"
    assemble_docx(minimal_selection, minimal_profile, _template(), out)

    with zipfile.ZipFile(out) as z:
        doc_xml = z.read("word/document.xml").decode()
        rels = z.read("word/_rels/document.xml.rels").decode()
    used = set(re.findall(r'r:id="([^"]+)"', doc_xml))
    have = set(re.findall(r'Id="([^"]+)"', rels))
    assert not (used - have), f"dangling relationship ids: {sorted(used - have)}"


def test_a_work_entry_still_shows_its_dates(
    minimal_profile, minimal_selection, tmp_path
):
    doc = _assemble(minimal_profile, minimal_selection, tmp_path)
    line = next(p for p in doc.paragraphs if p.text.startswith("Backend Engineer at"))
    assert line.text.endswith("January 2024 to current")
    assert line._p.find(qn("w:hyperlink")) is None


def test_frozen_prefix_unchanged(minimal_profile, minimal_selection, tmp_path):
    """Hard rules #9 and #10: nothing before the work heading may move.

    The frozen region is now a PREFIX — name, contact, citizenship AND the whole
    Education & Certificates block — because the Headless template puts education
    at the top instead of the bottom.
    """
    from src.endpoint.assembler import _frozen_canonical, _tailored_start

    template = Document(str(_template()))
    start = _tailored_start(template.element.body)
    before = _frozen_canonical(template.element.body, start)

    doc = _assemble(minimal_profile, minimal_selection, tmp_path)
    after = _frozen_canonical(doc.element.body, _tailored_start(doc.element.body))

    assert after == before
    assert start >= 6, "the frozen prefix must include the education lines"


def test_cloned_bullets_keep_their_numbering_and_entry_lines_their_tabs(
    minimal_profile, minimal_selection, tmp_path
):
    """The silent failure mode: a DOCX that opens fine and looks wrong.

    Asserting the numId VALUE, not merely that some numPr exists — a bullet
    cloned with the education numId would render with the wrong glyph and indent
    while passing a presence check.
    """
    from src.config import settings

    doc = _assemble(minimal_profile, minimal_selection, tmp_path)
    bullet = next(p for p in doc.paragraphs if p.text.startswith("Built APIs"))
    num_id = bullet._p.find(".//" + qn("w:numId"))
    assert num_id is not None, "bullet lost its <w:numPr> (no bullet glyph)"
    assert num_id.get(qn("w:val")) == str(settings.endpoint.render.bullet_num_id)

    line = next(p for p in doc.paragraphs if p.text.startswith("Backend Engineer at"))
    tabs = line._p.find(qn("w:pPr")).find(qn("w:tabs"))
    assert tabs is not None, "entry line lost its tab stops; dates would not align"


def test_a_cloned_bullet_keeps_the_prototype_formatting(
    minimal_profile, minimal_selection, tmp_path
):
    """Text differs; every formatting property must not."""
    import re

    from src.endpoint.assembler import _is_entry_bullet

    def _props(paragraph) -> str:
        """Canonical <w:pPr>, minus namespace declarations.

        python-docx rewrites the root nsmap on save, so the declarations carried
        on a c14n-serialised subtree differ between the template and the output
        even when every formatting property is identical. Stripping them keeps
        the assertion about formatting, which is what can actually break.
        """
        xml = etree.tostring(paragraph._p.find(qn("w:pPr")), method="c14n").decode()
        return re.sub(r'\sxmlns(:\w+)?="[^"]*"', "", xml)

    template = Document(str(_template()))
    proto = next(p for p in template.paragraphs if _is_entry_bullet(p._p))

    doc = _assemble(minimal_profile, minimal_selection, tmp_path)
    bullet = next(p for p in doc.paragraphs if p.text.startswith("Built APIs"))
    assert _props(bullet) == _props(proto)


def test_a_missing_bullet_is_an_error_not_a_blank_line(
    minimal_profile, minimal_selection, tmp_path
):
    from src.endpoint.assembler import AssemblerError, assemble_docx

    minimal_selection.entries[0].bullet_ids = ["b1", "nonexistent"]
    with pytest.raises(AssemblerError, match="nonexistent"):
        assemble_docx(minimal_selection, minimal_profile, _template(), tmp_path / "o.docx")


def test_education_over_the_cap_is_rejected(minimal_profile, minimal_selection, tmp_path):
    """The static region can only drift when the operator edits the template —
    exactly when nobody is checking."""
    from src.endpoint.assembler import AssemblerError, assemble_docx, _tailored_start
    import copy as _copy

    doc = Document(str(_template()))
    body = doc.element.body
    start = _tailored_start(body)
    edu = [c for c in list(body)[:start]
           if c.tag == qn("w:p") and c.find(".//" + qn("w:numId")) is not None]
    for _ in range(2):  # push it over the 3-line cap
        body.insert(start - 1, _copy.deepcopy(edu[0]))
    fat = tmp_path / "fat.docx"
    doc.save(str(fat))

    with pytest.raises(AssemblerError, match="caps it at 3"):
        assemble_docx(minimal_selection, minimal_profile, fat, tmp_path / "o.docx")


# ---------------------------------------------------------------------------
# Cache
# ---------------------------------------------------------------------------


def test_a_pre_pivot_selection_is_refused_rather_than_rendered():
    """v1 rows reference a profile that no longer exists and describe sections the
    template does not have. Rendering one would produce a resume that is not the
    one the operator was notified about, so it must fail loudly."""
    from src.endpoint.cache import StaleSelectionError, _load_selection

    v1 = {
        "job_id": "old1", "summary_id": "data_s1",
        "experiences": [{"exp_id": "e1", "title_alias": "Data Engineer",
                         "bullet_ids": ["market_data_b1"]}],
        "projects": [], "skills": {"categories": [], "familiar_with": []},
        "section_order": ["Work", "Skills", "Projects"],
        "cover_letter_text": "", "template_version": "deadbeef",
        "built_at": "2026-01-01T00:00:00Z",
    }
    with pytest.raises(StaleSelectionError, match="pre-pivot"):
        _load_selection(v1, "old1")


def test_a_current_selection_loads(minimal_selection):
    from src.endpoint.cache import _load_selection

    loaded = _load_selection(minimal_selection.model_dump(), "job001")
    assert loaded.version == 2
    assert len(loaded.entries) == 2


# ---------------------------------------------------------------------------
# FastAPI app
# ---------------------------------------------------------------------------


def test_app_404_on_unknown_job():
    """GET /resume/unknown.docx returns 404."""
    from fastapi.testclient import TestClient
    from src.endpoint.app import app

    client = TestClient(app, raise_server_exceptions=False)

    with patch("src.endpoint.app.session_scope") as mock_scope:
        mock_session = MagicMock()
        mock_scope.return_value.__enter__ = MagicMock(return_value=mock_session)
        mock_scope.return_value.__exit__ = MagicMock(return_value=False)

        with patch("src.endpoint.app.get_or_build", side_effect=KeyError("not found")):
            resp = client.get("/resume/unknown_job.docx")
            assert resp.status_code == 404


def test_app_400_on_bad_extension():
    """GET /resume/job123.txt returns 400."""
    from fastapi.testclient import TestClient
    from src.endpoint.app import app

    client = TestClient(app)
    resp = client.get("/resume/job123.txt")
    assert resp.status_code == 400


def test_app_health():
    """GET /health returns 200."""
    from fastapi.testclient import TestClient
    from src.endpoint.app import app

    client = TestClient(app)
    resp = client.get("/health")
    assert resp.status_code == 200
    assert resp.json() == {"status": "ok"}
