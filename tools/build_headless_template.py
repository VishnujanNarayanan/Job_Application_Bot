"""One-time: turn the pristine Headless template into the operator's copy.

This is Stage 0 of PIVOT_V3.md, done in code rather than by hand in Word. It
writes resumes/templates/headless_v1.docx, which is GITIGNORED -- the operator's
name, phone, email and links live in that asset, never in source (hard rule #21).

The assembler never touches any of what this writes: header and Education sit in
the frozen prefix, and are diff-asserted byte-identical on every render
(hard rules #9/#10).

Formatting rules come from resume guide/resume_method.md:
  - name font 14 bold, contact + location font 12, section body 10.5
  - Arial throughout, black and white EXCEPT phone/email/LinkedIn/portfolio
  - centre-aligned header, single column
  - Education & Certificates: no longer than 3 lines
Every run's existing rPr is reused, so sizes/fonts/spacing come from the
template itself rather than being asserted here.
"""
import copy
import os
import re
import sys
from pathlib import Path

import yaml

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from src.endpoint.assembler import apply_link_style  # noqa: E402
SRC = Path(os.environ.get(
    "HEADLESS_TEMPLATE",
    "/home/vishnu/projects/resume guide/Headless+Resume+Template.docx",
))
OUT = ROOT / "resumes" / "templates" / "headless_v1.docx"

# Everything identifying the operator is READ, never written here: hard rule #21
# forbids an operator-specific literal in source. The OUTPUT file carries that
# content, which is why resumes/templates/ is gitignored.
_profile = yaml.safe_load((ROOT / "master_profile.yaml").read_text())
_p = _profile["personal"]
# Hand-trimmed display strings, if the operator set them. Deriving from the
# profile yields the FULL degree title and full course names, which overrun both
# the 3-line section cap and the ~89-character entry line.
_hdr = dict(yaml.safe_load((ROOT / "config" / "config.yaml").read_text())
            .get("operator", {}).get("resume_header") or {})

NAME = _p["name"]
PHONE = _p.get("phone", "")
EMAIL = _p.get("email", "")
LOCATION = _p.get("location", "")
LINKS = {
    "tel": "tel://" + re.sub(r"\D", "", PHONE)[-10:] + "/",
    "mail": f"mailto:{EMAIL}",
    # The guide's slot is "Phone | Email | LinkedIn/Portfolio", so a personal
    # site takes precedence over a code host when both are present.
    "portfolio": _hdr.get("portfolio_url") or _p.get("portfolio") or _p.get("github", ""),
    "linkedin": _p.get("linkedin", ""),
    "certs": _p.get("certificates_link", ""),
}


def _education_line() -> tuple[str, str]:
    """Degree + institution, plus the right-hand status. First entry only.

    The method caps this section at three lines, so school results are dropped:
    once a degree exists they are noise competing for the budget.
    """
    edu = (_profile.get("education") or [{}])[0]
    degree = edu.get("degree", "")
    inst = edu.get("institution", "").split("(")[0].strip().rstrip(",")
    return f"{degree}, {inst}".strip(", "), "Status - Graduated"


def _certificates_line() -> str:
    """Every certificate on ONE comma-joined line, issuer parentheses stripped."""
    names = []
    for c in _profile.get("certifications") or []:
        n = c.get("name", "")
        n = n.split("\u2014")[-1].strip() if "\u2014" in n else n
        n = re.sub(r"\s*\([^)]*\)\s*$", "", n).strip()
        if n and n not in names:
            names.append(n)
    return ", ".join(names)


_derived_edu, _derived_status = _education_line()
EDU = _hdr.get("education_line") or _derived_edu
EDU_RIGHT = _hdr.get("education_status") or _derived_status
CERTS = _hdr.get("certificates_line") or _certificates_line()

doc = Document(str(SRC))
body = doc.element.body
paras = doc.paragraphs

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"


def rels_part():
    return doc.part


def add_hyperlink_rel(url: str) -> str:
    """Register an external relationship and return its rId."""
    return doc.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )


def styled_run(proto_run, text: str, *, blue: bool, underline: bool):
    """Clone a template run, keeping its rPr, and set text + link styling.

    Colour and underline are the ONLY properties touched: the guide permits blue
    for phone/email/portfolio links and nothing else, and everything about font,
    size and spacing must come from the template.
    """
    r = copy.deepcopy(proto_run)
    for t in r.findall(qn("w:t")):
        r.remove(t)
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    # Only touch colour/underline when ADDING them. Stripping them from a plain
    # run drops the template's explicit <w:color w:val="000000"/>, and the run
    # then inherits its style's colour -- which turned the name blue, because
    # Heading 1's own colour is blue.
    if blue or underline:
        for tag in ("w:color", "w:u"):
            found = rPr.find(qn(tag))
            if found is not None:
                rPr.remove(found)
    if blue:
        c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1"); rPr.append(c)
    if underline:
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    t = OxmlElement("w:t")
    t.text = text
    if text != text.strip():
        t.set(XMLSPACE, "preserve")
    r.append(t)
    return r


def set_paragraph(p, pieces):
    """Rebuild a paragraph from (text, url|None) pieces, reusing run formatting."""
    proto = p.runs[0]._r if p.runs else None
    if proto is None:
        return
    sep_proto = None
    for r in p.runs:                       # keep the Symbol-font separator run
        if (r.text or "").strip() == "⎪":
            sep_proto = r._r
            break
    for r in list(p._p.findall(qn("w:r"))):
        p._p.remove(r)
    for hl in list(p._p.findall(qn("w:hyperlink"))):
        p._p.remove(hl)

    for text, url in pieces:
        if text == "⎪" and sep_proto is not None:
            p._p.append(styled_run(sep_proto, " ⎪ ", blue=False, underline=False))
            continue
        if url:
            rid = add_hyperlink_rel(url)
            hl = OxmlElement("w:hyperlink")
            hl.set(qn("r:id"), rid)
            # Word's own followed-link marker. It does NOT decide whether the
            # PDF gets a link annotation -- the run's rStyle does, which is why
            # apply_link_style is called below. See its docstring.
            hl.set(qn("w:history"), "1")
            run = styled_run(proto, text, blue=True, underline=True)
            apply_link_style(doc, run)
            hl.append(run)
            p._p.append(hl)
        else:
            p._p.append(styled_run(proto, text, blue=False, underline=False))


# --- paragraph 0: the name -------------------------------------------------
set_paragraph(paras[0], [(NAME, None)])

# --- paragraph 1: contact line ---------------------------------------------
set_paragraph(paras[1], [
    (PHONE, LINKS["tel"]), ("⎪", None),
    (EMAIL, LINKS["mail"]), ("⎪", None),
    ("Portfolio", LINKS["portfolio"]), ("⎪", None),
    ("LinkedIn", LINKS["linkedin"]), ("⎪", None),
    ("Certificates", LINKS["certs"]),
])

# --- paragraph 2: location --------------------------------------------------
# The guide's slot is "Citizenship status at City, State". Indian market, Indian
# roles: there is no visa question to answer, so the line carries the location
# only rather than stating something a recruiter here never asks.
set_paragraph(paras[2], [(LOCATION, None)])

# --- paragraphs 6-7: Education & Certificates (max 3 lines) ----------------
# Only the degree earns its own line. Class X/XII are dropped: the guide caps
# this section at three lines and school results are noise once a degree exists.
edu_p, cert_p = paras[6], paras[7]

def set_tabbed(p, left: str, right: str):
    runs = p._p.findall(qn("w:r"))
    tab_i = next((i for i, r in enumerate(runs) if r.find(qn("w:tab")) is not None), -1)
    if tab_i < 0:
        return
    pre, post = runs[:tab_i], runs[tab_i + 1:]
    def put(rs, text):
        if not rs:
            return
        t = rs[0].find(qn("w:t"))
        if t is None:
            t = OxmlElement("w:t"); rs[0].append(t)
        t.text = text
        if text != text.strip():
            t.set(XMLSPACE, "preserve")
        for r in rs[1:]:
            tt = r.find(qn("w:t"))
            if tt is not None:
                tt.text = ""
    put(pre, left); put(post, right)

set_tabbed(edu_p, EDU, EDU_RIGHT)
set_tabbed(cert_p, CERTS, "")

# --- bullet glyph ----------------------------------------------------------
# The template defines the glyph as U+25CF BLACK CIRCLE in "Noto Sans Symbols",
# with size inheriting from the 10.5pt body text. Word does not have that font
# so it substitutes a small dot; LibreOffice on Linux DOES have it and draws the
# true circle, so the PDF a recruiter receives had visibly heavier bullets than
# the document the operator designed. U+2022 BULLET in Arial renders the same
# conventional size in both, which is what the method's "black and white, plain"
# formatting asks for.
numbering = doc.part.numbering_part.element
# U+2022 BULLET in Arial at 16pt. The template defines U+25CF BLACK CIRCLE in
# "Noto Sans Symbols" at the inherited body size; Word lacks that font and
# substitutes a small dot, while LibreOffice has it and draws a heavy circle, so
# the DOCX and the PDF disagreed. A bullet pinned to an explicit size renders
# identically in both. 16pt was chosen by rendering a ladder of sizes and
# picking by eye against the guide's own Example Resume.
BULLET, GLYPH_FONT, GLYPH_HALF_POINTS = "\u2022", "Arial", 32
changed = 0
for lvl in numbering.iter(qn("w:lvl")):
    text = lvl.find(qn("w:lvlText"))
    if text is None or text.get(qn("w:val")) != "\u25cf":
        continue
    text.set(qn("w:val"), BULLET)
    rPr = lvl.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); lvl.append(rPr)
    fonts = rPr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts"); rPr.insert(0, fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attr), GLYPH_FONT)
    for tag in ("w:sz", "w:szCs"):
        existing = rPr.find(qn(tag))
        if existing is not None:
            rPr.remove(existing)
        size = OxmlElement(tag)
        size.set(qn("w:val"), str(GLYPH_HALF_POINTS))
        rPr.append(size)
    changed += 1

# --- section headings become real headings ---------------------------------
# The template ships them as Normal paragraphs that are merely bolded, so they
# carry no outline level: they do not collapse in Word, do not appear in the
# navigation pane, and give a document parser no structure to read. Promoting
# them to Heading 2 fixes all three.
#
# Appearance is unaffected. Heading 2's own style is Play 16pt blue, but every
# run here carries Arial / 10.5 / bold / black as DIRECT formatting, and direct
# formatting beats a style. Spacing is likewise pinned on the paragraph itself.
SECTION_HEADING_STYLE = "Heading 2"
for idx in (5, 9):
    para = paras[idx]
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para._p.insert(0, pPr)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), "Heading2")
    # Pin spacing so the heading style's own before/after cannot shift the layout.
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    # Pin colour explicitly. As a Normal paragraph the heading inherited black;
    # as Heading 2 it would inherit that style's 0f4761 blue instead, which is
    # the same trap that turned the name blue. The method allows colour only on
    # phone/email/portfolio links, so a blue section heading is a real defect.
    for run in para._p.findall(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            run.insert(0, rPr)
        existing = rPr.find(qn("w:color"))
        if existing is not None:
            rPr.remove(existing)
        colour = OxmlElement("w:color")
        colour.set(qn("w:val"), "000000")
        rPr.append(colour)
print(f"promoted 2 section headings to {SECTION_HEADING_STYLE}")


# --- pin the line box on every bulleted paragraph --------------------------
# With lineRule="auto" a line grows to fit its tallest content, so a 16pt bullet
# glyph sitting in 10.5pt text stretched the spacing from a uniform 18.1pt to an
# uneven 19-23pt. Measured across sizes: ANY glyph above 10.5pt inflates it, and
# progressively. "exact" fixes the height at 360 twips (18pt), the value 1.5
# spacing already resolves to, so the glyph can be sized for legibility without
# moving the lines.
for _p in doc.paragraphs:
    _pPr = _p._p.find(qn("w:pPr"))
    if _pPr is None or _pPr.find(qn("w:numPr")) is None:
        continue
    _sp = _pPr.find(qn("w:spacing"))
    if _sp is None:
        _sp = OxmlElement("w:spacing")
        _pPr.append(_sp)
    _sp.set(qn("w:line"), "360")
    _sp.set(qn("w:lineRule"), "exact")


# --- one blank line before every section heading ---------------------------
# The template ships TWO empty paragraphs between the contact block and
# "Education & Certificates" but only ONE before the work heading: 36pt against
# 18pt at the same kind of boundary. The doubled gap is what made the header
# look detached from the body. Collapse any run of blanks in the frozen prefix
# to a single paragraph so both boundaries breathe the same amount.
_body = doc.element.body
_children = list(_body)
_run = []
for _c in _children:
    _is_blank = (_c.tag == qn("w:p")
                 and not "".join(t.text or "" for t in _c.iter(qn("w:t"))).strip()
                 and _c.find(qn("w:pPr")) is not None
                 and _c.find(qn("w:pPr")).find(qn("w:numPr")) is None)
    if _is_blank:
        _run.append(_c)
        continue
    for _extra in _run[1:]:
        _body.remove(_extra)
    _run = []
for _extra in _run[1:]:
    _body.remove(_extra)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"bullet glyph normalised on {changed} level(s)")
print(f"wrote {OUT}  ({OUT.stat().st_size:,} bytes)")
